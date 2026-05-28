"""
Format existing docx files according to a template preset.

Applies:
- Page setup (size, margins)
- Default font and paragraph style
- Heading styles
- Header/footer (page number)
- Optional: regenerate TOC
"""

from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .config import TemplateConfig
from .converter.docx_builder import _setup_page, _setup_styles, _setup_header_footer
from .converter.toc import insert_toc


def format_existing_docx(input_path: str, output_path: str, config: TemplateConfig,
                         add_toc: bool = False):
    """Load an existing docx and apply template formatting."""
    doc = Document(input_path)

    # 1. Page setup
    _setup_page(doc, config)

    # 2. Styles
    _setup_styles(doc, config)

    # 3. Header / footer
    _setup_header_footer(doc, config)

    # 4. Normalize paragraph fonts for body paragraphs
    for para in doc.paragraphs:
        _normalize_paragraph(para, config)

    # 5. Optionally insert TOC at beginning
    if add_toc and doc.paragraphs:
        first_p = doc.paragraphs[0]
        # Insert before first paragraph
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run("目录")
        run.bold = True
        run.font.size = Pt(config.heading.h1_size_pt)
        _set_east_asia(run, config.font.heading_chinese)
        toc_p = doc.add_paragraph()
        toc_p.paragraph_format.first_line_indent = Cm(0)
        insert_toc(toc_p, "目录")
        # Move to top
        body = doc.element.body
        body.insert(0, toc_p._element)
        body.insert(0, p._element)

    doc.save(output_path)


def _normalize_paragraph(para, config: TemplateConfig):
    """Ensure paragraph runs use template fonts."""
    for run in para.runs:
        if not run.font.name:
            run.font.name = config.font.english
        run.font.size = Pt(config.body.size_pt)
        _set_east_asia(run, config.font.chinese)


def _set_east_asia(run, font_name: str):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
