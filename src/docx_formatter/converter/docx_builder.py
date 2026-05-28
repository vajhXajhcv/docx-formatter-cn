"""
Build a python-docx Document from parsed Markdown AST.
"""

import os
import re
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.shared import Pt, Mm, Cm, Inches, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from latex2mathml.converter import convert as latex_to_mathml

from .markdown_parser import (
    BlockElement, BlockHeading, BlockParagraph, BlockFormula,
    BlockTable, BlockCode, BlockQuote, BlockListItem, BlockImage, BlockBlank, BlockPageBreak,
    InlineElement, InlineText, InlineBold, InlineItalic, InlineUnderline, InlineStrikethrough, InlineFormula,
    InlineCitation, InlineImage, InlineLink,
)
from ..config import TemplateConfig
from ..formula.mathml_to_omml import mathml_to_omml_element
from .toc import insert_toc


def build_docx(blocks: List[BlockElement], config: TemplateConfig, output_path: str,
               image_base_dir: Optional[str] = None):
    """Build and save a docx from parsed Markdown blocks."""
    doc = Document()
    _setup_page(doc, config)
    _setup_styles(doc, config)
    _setup_header_footer(doc, config)

    # State
    chapter_counter = [0]
    formula_counter = [0]
    table_counter = [0]
    figure_counter = [0]

    ctx = {
        "doc": doc,
        "config": config,
        "image_base_dir": image_base_dir or "",
        "chapter_counter": chapter_counter,
        "formula_counter": formula_counter,
        "table_counter": table_counter,
        "figure_counter": figure_counter,
    }

    for block in blocks:
        _render_block(block, ctx)

    doc.save(output_path)


def _setup_page(doc: Document, config: TemplateConfig):
    section = doc.sections[0]
    pc = config.page
    section.page_width = Mm(pc.width_mm)
    section.page_height = Mm(pc.height_mm)
    section.top_margin = Mm(pc.margin_top_mm)
    section.bottom_margin = Mm(pc.margin_bottom_mm)
    section.left_margin = Mm(pc.margin_left_mm)
    section.right_margin = Mm(pc.margin_right_mm)
    section.header_distance = Mm(pc.header_distance_mm)
    section.footer_distance = Mm(pc.footer_distance_mm)


def _setup_styles(doc: Document, config: TemplateConfig):
    """Configure default document font and paragraph style."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = config.font.english
    font.size = Pt(config.body.size_pt)
    # East Asian font
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), config.font.chinese)
    rfonts.set(qn("w:ascii"), config.font.english)
    rfonts.set(qn("w:hAnsi"), config.font.english)

    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = config.body.line_spacing
    pf.alignment = _align_wd(config.body.alignment)
    pf.first_line_indent = Cm(config.body.first_line_indent_chars * 0.35)

    # Heading styles
    for i, (size, align) in enumerate([
        (config.heading.h1_size_pt, config.heading.alignment),
        (config.heading.h2_size_pt, config.heading.alignment),
        (config.heading.h3_size_pt, config.heading.alignment),
        (config.heading.h4_size_pt, config.heading.alignment),
    ], start=1):
        try:
            hs = doc.styles[f"Heading {i}"]
        except KeyError:
            continue
        hf = hs.font
        hf.name = config.font.heading_english
        hf.size = Pt(size)
        hf.bold = config.heading.bold
        h_rpr = hs.element.get_or_add_rPr()
        h_rfonts = h_rpr.find(qn("w:rFonts"))
        if h_rfonts is None:
            h_rfonts = OxmlElement("w:rFonts")
            h_rpr.insert(0, h_rfonts)
        h_rfonts.set(qn("w:eastAsia"), config.font.heading_chinese)
        h_rfonts.set(qn("w:ascii"), config.font.heading_english)
        h_rfonts.set(qn("w:hAnsi"), config.font.heading_english)
        hpf = hs.paragraph_format
        hpf.alignment = _align_wd(align)
        hpf.first_line_indent = Cm(0)


def _setup_header_footer(doc: Document, config: TemplateConfig):
    section = doc.sections[0]
    if config.header_text:
        header = section.header
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run(config.header_text)
        run.font.size = Pt(9)
        run.font.name = config.font.english
        _set_east_asia_font(run, config.font.chinese)

    if config.footer_page_number:
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Insert page number field
        run = fp.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        run._element.append(fldChar1)

        run2 = fp.add_run()
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = " PAGE "
        run2._element.append(instrText)

        run3 = fp.add_run()
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run3._element.append(fldChar2)


def _render_block(block: BlockElement, ctx: dict):
    doc = ctx["doc"]
    config = ctx["config"]

    if isinstance(block, BlockHeading):
        # Detect TOC heading
        heading_text = "".join(
            e.text if isinstance(e, InlineText) else
            e.latex if isinstance(e, InlineFormula) else str(e)
            for e in block.inline
        ).strip()
        if heading_text in ("目录", "目次", "Table of Contents", "Contents") and config.toc_enabled:
            para = doc.add_paragraph()
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Cm(0)
            run = para.add_run(heading_text)
            run.bold = True
            run.font.size = Pt(config.heading.h1_size_pt)
            run.font.name = config.font.heading_english
            _set_east_asia_font(run, config.font.heading_chinese)
            # Insert TOC field
            toc_para = doc.add_paragraph()
            toc_para.paragraph_format.first_line_indent = Cm(0)
            insert_toc(toc_para, heading_text)
            return
        para = doc.add_paragraph()
        try:
            para.style = f"Heading {min(block.level, 4)}"
        except KeyError:
            pass
        _render_inline(block.inline, para, ctx)
        if block.level == 1:
            ctx["chapter_counter"][0] += 1
            # Reset sub-counters if desired
    elif isinstance(block, BlockParagraph):
        para = doc.add_paragraph()
        _render_inline(block.inline, para, ctx)
    elif isinstance(block, BlockFormula):
        _render_formula_block(block, ctx)
    elif isinstance(block, BlockTable):
        _render_table(block, ctx)
    elif isinstance(block, BlockCode):
        for line in block.lines:
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.left_indent = Cm(1.0)
            run = para.add_run(line)
            run.font.name = config.font.code
            run.font.size = Pt(10)
            _set_no_proof(run)
    elif isinstance(block, BlockQuote):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.5 * block.level)
        _render_inline(block.inline, para, ctx)
        # Add left border to paragraph
        pPr = para._element.get_or_add_pPr()
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is None:
            pBdr = OxmlElement("w:pBdr")
            pPr.append(pBdr)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "24")
        left.set(qn("w:space"), "4")
        left.set(qn("w:color"), "CCCCCC")
        pBdr.append(left)
        # Light gray shading
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "F5F5F5")
        pPr.append(shd)
    elif isinstance(block, BlockListItem):
        para = doc.add_paragraph()
        marker = block.marker
        if block.ordered:
            prefix = f"{marker} "
        else:
            prefix = f"{marker} "
        run = para.add_run(prefix)
        _render_inline(block.inline, para, ctx, skip_first_space=True)
    elif isinstance(block, BlockImage):
        _render_image(block, ctx)
    elif isinstance(block, BlockBlank):
        doc.add_paragraph()
    elif isinstance(block, BlockPageBreak):
        para = doc.add_paragraph()
        run = para.add_run()
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        run._element.append(br)


def _render_inline(inline: List[InlineElement], para, ctx: dict, skip_first_space: bool = False):
    config = ctx["config"]
    for elem in inline:
        if isinstance(elem, InlineText):
            text = elem.text
            if skip_first_space and text.startswith(" "):
                text = text[1:]
                skip_first_space = False
            if text:
                run = para.add_run(text)
                _apply_body_font(run, config)
        elif isinstance(elem, InlineBold):
            run = para.add_run(elem.text)
            run.bold = True
            _apply_body_font(run, config)
        elif isinstance(elem, InlineItalic):
            run = para.add_run(elem.text)
            run.italic = True
            _apply_body_font(run, config)
        elif isinstance(elem, InlineUnderline):
            run = para.add_run(elem.text)
            run.underline = True
            _apply_body_font(run, config)
        elif isinstance(elem, InlineStrikethrough):
            run = para.add_run(elem.text)
            run.font.strike = True
            _apply_body_font(run, config)
        elif isinstance(elem, InlineFormula):
            _insert_inline_formula(para, elem.latex)
        elif isinstance(elem, InlineCitation):
            run = para.add_run(f"[{elem.number}]")
            run.font.superscript = True
            _apply_body_font(run, config)
        elif isinstance(elem, InlineImage):
            _render_inline_image(para, elem, ctx)
        elif isinstance(elem, InlineLink):
            run = para.add_run(elem.text)
            run.font.color.rgb = None  # hyperlink blue handled by Word if we add hyperlink properly
            _apply_body_font(run, config)
            # Full hyperlink support would require adding a relationship; skip for MVP


def _apply_body_font(run, config: TemplateConfig):
    run.font.name = config.font.english
    run.font.size = Pt(config.body.size_pt)
    _set_east_asia_font(run, config.font.chinese)


def _set_east_asia_font(run, font_name: str):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def _insert_inline_formula(paragraph, latex: str):
    try:
        mathml = latex_to_mathml(latex)
        omath = mathml_to_omml_element(mathml)
        paragraph._element.append(omath)
    except Exception:
        # Fallback: plain text
        run = paragraph.add_run(f"${latex}$")
        run.font.name = "Cambria Math"
        run.font.italic = True


def _render_formula_block(block: BlockFormula, ctx: dict):
    doc = ctx["doc"]
    config = ctx["config"]
    ctx["formula_counter"][0] += 1
    number = block.number or str(ctx["formula_counter"][0])

    # Use a 3-column borderless table: left margin | formula | number
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table, top=False, bottom=False, left=False, right=False, inside_h=False, inside_v=False)

    # Column widths: ~1cm | rest | ~1cm
    total_width = Mm(config.page.width_mm - config.page.margin_left_mm - config.page.margin_right_mm)
    margin_width = Mm(10)
    formula_width = total_width - margin_width * 2

    # Set column widths via tblGrid
    tblGrid = table._element.find(qn("w:tblGrid"))
    if tblGrid is not None:
        for i, w in enumerate([margin_width, formula_width, margin_width]):
            gridCol = tblGrid.findall(qn("w:gridCol"))[i]
            gridCol.set(qn("w:w"), str(int(w)))

    row = table.rows[0]
    left_cell = row.cells[0]
    formula_cell = row.cells[1]
    right_cell = row.cells[2]

    # Left empty
    left_cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Center formula
    fp = formula_cell.paragraphs[0]
    fp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.first_line_indent = Cm(0)
    try:
        mathml = latex_to_mathml(block.latex)
        omath = mathml_to_omml_element(mathml)
        fp._element.append(omath)
    except Exception:
        run = fp.add_run(block.latex)
        run.font.name = "Cambria Math"
        run.font.italic = True

    # Right number
    rp = right_cell.paragraphs[0]
    rp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.first_line_indent = Cm(0)
    run = rp.add_run(f"({number})")
    run.font.name = config.font.english
    run.font.size = Pt(config.body.size_pt)
    _set_east_asia_font(run, config.font.chinese)


def _render_table(block: BlockTable, ctx: dict):
    doc = ctx["doc"]
    config = ctx["config"]
    ctx["table_counter"][0] += 1

    num_cols = len(block.headers)
    num_rows = 1 + len(block.rows)
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Three-line table borders
    tc = config.table
    if tc.three_line:
        _set_table_borders(
            table,
            top=True, top_width_pt=tc.header_border_width_pt,
            bottom=True, bottom_width_pt=tc.bottom_border_width_pt,
            left=False, right=False, inside_h=False, inside_v=False,
        )
    else:
        _set_table_borders(table, top=True, bottom=True, left=True, right=True, inside_h=True, inside_v=True)

    # Header row
    for j, header in enumerate(block.headers):
        cell = table.rows[0].cells[j]
        p = cell.paragraphs[0]
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        from .markdown_parser import _parse_inline
        inline_elems = _parse_inline(header)
        for elem in inline_elems:
            if isinstance(elem, InlineText):
                run = p.add_run(elem.text)
                run.bold = True
            elif isinstance(elem, InlineBold):
                run = p.add_run(elem.text)
                run.bold = True
            elif isinstance(elem, InlineItalic):
                run = p.add_run(elem.text)
                run.bold = True
                run.italic = True
            elif isinstance(elem, InlineFormula):
                _insert_inline_formula(p, elem.latex)
            else:
                run = p.add_run(str(elem))
                run.bold = True
            run.font.size = Pt(tc.font_size_pt)
            run.font.name = config.font.english
            _set_east_asia_font(run, config.font.chinese)
        # Bottom border for header row (thin line)
        if tc.three_line:
            _set_cell_border(cell, bottom=True, width_pt=tc.header_border_width_pt / 2 or 0.75)

    # Data rows
    for i, row_data in enumerate(block.rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            p = cell.paragraphs[0]
            align = _align_wd(block.alignments[j]) if j < len(block.alignments) else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.alignment = align
            p.paragraph_format.first_line_indent = Cm(0)
            # Parse inline elements in cell text (formulas, bold, etc.)
            from .markdown_parser import _parse_inline
            inline_elems = _parse_inline(cell_text)
            _render_inline(inline_elems, p, ctx)
            # Last row bottom border (thick)
            if tc.three_line and i == len(block.rows) - 1:
                _set_cell_border(cell, bottom=True, width_pt=tc.bottom_border_width_pt)

    # Caption
    caption_para = doc.add_paragraph()
    caption_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_para.paragraph_format.first_line_indent = Cm(0)
    prefix = config.caption.table_prefix
    caption_text = f"{prefix} {ctx['chapter_counter'][0]}-{ctx['table_counter'][0]}"
    run = caption_para.add_run(caption_text)
    run.bold = config.caption.bold
    run.font.size = Pt(config.caption.font_size_pt)
    run.font.name = config.font.english
    _set_east_asia_font(run, config.font.chinese)


def _resolve_image_path(url: str, base_dir: str) -> Path:
    """Resolve image path, downloading from network if necessary."""
    from urllib.parse import urlparse
    from urllib.request import urlopen
    import tempfile

    parsed = urlparse(url)
    if parsed.scheme in ("http", "https"):
        # Download to temp file
        suffix = Path(parsed.path).suffix or ".jpg"
        fd, tmp = tempfile.mkstemp(suffix=suffix)
        os.close(fd)
        try:
            with urlopen(url, timeout=30) as resp:
                with open(tmp, "wb") as f:
                    f.write(resp.read())
            return Path(tmp)
        except Exception:
            Path(tmp).unlink(missing_ok=True)
            raise
    else:
        base = Path(base_dir) if base_dir else Path(".")
        p = base / url
        if p.exists():
            return p
        raise FileNotFoundError(f"Image not found: {p}")


def _render_image(block: BlockImage, ctx: dict):
    doc = ctx["doc"]
    config = ctx["config"]
    ctx["figure_counter"][0] += 1

    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Cm(0)
    try:
        img_path = _resolve_image_path(block.url, ctx.get("image_base_dir") or "")
        run = para.add_run()
        kwargs = {}
        if block.width:
            kwargs["width"] = Inches(block.width / 72)
        elif block.height:
            kwargs["height"] = Inches(block.height / 72)
        else:
            kwargs["width"] = Inches(4.0)
        run.add_picture(str(img_path), **kwargs)
    except Exception as e:
        para.add_run(f"[Image error: {e}]")

    # Caption
    caption_para = doc.add_paragraph()
    caption_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_para.paragraph_format.first_line_indent = Cm(0)
    prefix = config.caption.figure_prefix
    caption_text = f"{prefix} {ctx['chapter_counter'][0]}-{ctx['figure_counter'][0]}"
    run = caption_para.add_run(caption_text)
    run.bold = config.caption.bold
    run.font.size = Pt(config.caption.font_size_pt)
    run.font.name = config.font.english
    _set_east_asia_font(run, config.font.chinese)


def _render_inline_image(paragraph, elem: InlineImage, ctx: dict):
    try:
        img_path = _resolve_image_path(elem.url, ctx.get("image_base_dir") or "")
        run = paragraph.add_run()
        kwargs = {}
        if elem.width:
            kwargs["width"] = Inches(elem.width / 72)
        elif elem.height:
            kwargs["height"] = Inches(elem.height / 72)
        else:
            kwargs["width"] = Inches(4.0)
        run.add_picture(str(img_path), **kwargs)
    except Exception as e:
        paragraph.add_run(f"[Image error: {e}]")


def _align_wd(align: str):
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(align, WD_ALIGN_PARAGRAPH.LEFT)


def _set_no_proof(run):
    rpr = run._element.get_or_add_rPr()
    no_proof = OxmlElement("w:noProof")
    rpr.append(no_proof)


def _set_table_borders(table, *, top=False, bottom=False, left=False, right=False,
                       inside_h=False, inside_v=False,
                       top_width_pt=1.5, bottom_width_pt=1.5, left_width_pt=0.5,
                       right_width_pt=0.5, inside_h_width_pt=0.5, inside_v_width_pt=0.5):
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is not None:
        tblPr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    tblPr.append(borders)

    def _add_border(name, show, width_pt):
        el = OxmlElement(f"w:{name}")
        if show:
            el.set(qn("w:val"), "single")
            # 1 pt = 8 eighth-points
            el.set(qn("w:sz"), str(int(width_pt * 8)))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "none")
            el.set(qn("w:sz"), "0")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "auto")
        borders.append(el)

    _add_border("top", top, top_width_pt)
    _add_border("bottom", bottom, bottom_width_pt)
    _add_border("left", left, left_width_pt)
    _add_border("right", right, right_width_pt)
    _add_border("insideH", inside_h, inside_h_width_pt)
    _add_border("insideV", inside_v, inside_v_width_pt)


def _set_cell_border(cell, *, top=None, bottom=None, left=None, right=None, width_pt=0.75):
    tc = cell._element
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is not None:
        tcPr.remove(borders)
    borders = OxmlElement("w:tcBorders")
    tcPr.append(borders)

    def _add(name, show):
        el = OxmlElement(f"w:{name}")
        if show:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(int(width_pt * 8)))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "none")
            el.set(qn("w:sz"), "0")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "auto")
        borders.append(el)

    _add("top", top if top is not None else False)
    _add("bottom", bottom if bottom is not None else False)
    _add("left", left if left is not None else False)
    _add("right", right if right is not None else False)
