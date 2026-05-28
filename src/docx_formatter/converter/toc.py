"""
Insert Table of Contents (TOC) field into a python-docx Document.
"""

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def insert_toc(paragraph, title: str = "目录"):
    """Insert a Word TOC field at the given paragraph."""
    run = paragraph.add_run()

    # Begin field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._element.append(fldChar1)

    # Instruction text
    run2 = paragraph.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._element.append(instrText)

    # Separate
    run3 = paragraph.add_run()
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run3._element.append(fldChar2)

    # Placeholder text
    run4 = paragraph.add_run(title)

    # End field
    run5 = paragraph.add_run()
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run5._element.append(fldChar3)
